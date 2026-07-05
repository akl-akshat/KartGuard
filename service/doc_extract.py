"""Policy-document text extraction — PDF, DOCX, Markdown, plain text.

Clients upload their returns policy as whatever their legal team produced. This module turns
the file into clean paragraph text for the RAG pipeline (``policy_store.chunk_document``
splits on blank lines, so extraction inserts them between structural units):

* **PDF** — PyPDF2 per-page text; pages separated by blank lines. Handles long, multi-page
  documents (page cap keeps a pathological upload from stalling the request).
* **DOCX** — python-docx paragraphs (incl. simple tables), blank-line separated, heading
  styles rendered as markdown ``##`` so the chunker keeps section context.
* **.md / .txt** — decoded as UTF-8 (best-effort fallback encodings).
"""

from __future__ import annotations

import io

MAX_PDF_PAGES = 300
MAX_CHARS = 400_000


class ExtractionError(ValueError):
    pass


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        text = _from_pdf(data)
    elif name.endswith(".docx"):
        text = _from_docx(data)
    elif name.endswith((".md", ".txt", ".text", ".markdown")):
        text = _decode(data)
    else:
        # sniff: PDFs and DOCX have unmistakable magic bytes
        if data[:5] == b"%PDF-":
            text = _from_pdf(data)
        elif data[:2] == b"PK":
            text = _from_docx(data)
        else:
            text = _decode(data)
    text = text.strip()
    if len(text) < 40:
        raise ExtractionError("could not extract readable policy text from the document")
    return text[:MAX_CHARS]


def _decode(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ExtractionError("unsupported text encoding")


def _from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        from PyPDF2 import PdfReader  # type: ignore
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"unreadable PDF: {exc.__class__.__name__}") from exc
    pages = []
    for i, page in enumerate(reader.pages):
        if i >= MAX_PDF_PAGES:
            break
        try:
            pages.append((page.extract_text() or "").strip())
        except Exception:  # noqa: BLE001 — a single bad page must not sink the document
            continue
    return "\n\n".join(_reparagraph(p) for p in pages if p)


def _reparagraph(page_text: str) -> str:
    """PDF extraction yields hard line breaks, not paragraphs. Rebuild them: a line that ends a
    sentence closes a paragraph; other line breaks are soft wraps and join with a space."""
    paras: list[list[str]] = [[]]
    for line in page_text.splitlines():
        line = line.strip()
        if not line:
            if paras[-1]:
                paras.append([])
            continue
        paras[-1].append(line)
        if line.endswith((".", "!", "?", ":", ";")):
            paras.append([])
    return "\n\n".join(" ".join(p) for p in paras if p)


def _from_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("DOCX support unavailable on this deployment") from exc
    try:
        d = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"unreadable DOCX: {exc.__class__.__name__}") from exc
    parts: list[str] = []
    for p in d.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        style = (p.style.name or "").lower() if p.style is not None else ""
        if style.startswith("heading"):
            t = "## " + t
        parts.append(t)
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" · ".join(cells))
    return "\n\n".join(parts)
